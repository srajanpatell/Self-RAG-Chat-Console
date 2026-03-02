from __future__ import annotations
from io import BytesIO
import uuid

from fastapi import Depends, FastAPI, File, HTTPException, UploadFile
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession

from .config import TOP_K
from .db import SessionLocal, init_db
from .embedder import Embedder
from .llm import LLMClient
from .models import Chunk, Document
from .retrieval import chunk_text, hybrid_search, upsert_chunk_fts
from .schemas import ChatRequest, ChatResponse, IngestTextRequest, Source

app = FastAPI(title="Self-RAG Chatbot")
embedder = Embedder()
llm = LLMClient()


async def get_db():
    async with SessionLocal() as session:
        yield session


@app.on_event("startup")
async def startup_event() -> None:
    await init_db()


@app.get("/health")
async def health() -> dict:
    return {"status": "ok"}


async def ingest_document(session: AsyncSession, title: str, content: str) -> dict:
    document = Document(id=str(uuid.uuid4()), title=title, content=content)
    session.add(document)

    chunks = chunk_text(content)
    chunk_ids: list[str] = []
    for idx, chunk in enumerate(chunks):
        vector = embedder.encode(chunk)
        chunk_id = str(uuid.uuid4())
        c = Chunk(
            id=chunk_id,
            document_id=document.id,
            chunk_index=idx,
            chunk_text=chunk,
            embedding=vector,
        )
        session.add(c)
        chunk_ids.append(chunk_id)

    await session.flush()

    for chunk_id in chunk_ids:
        await upsert_chunk_fts(session, chunk_id)

    await session.commit()
    return {"document_id": document.id, "title": title, "chunks": len(chunks)}


@app.post("/ingest/text")
async def ingest_text(payload: IngestTextRequest, session: AsyncSession = Depends(get_db)):
    return await ingest_document(session, payload.title, payload.text)


@app.post("/ingest/file")
async def ingest_file(file: UploadFile = File(...), session: AsyncSession = Depends(get_db)):
    raw = await file.read()
    ext = (file.filename or "").lower().split(".")[-1]

    if ext == "pdf":
        pdf = PdfReader(BytesIO(raw))
        text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    elif ext == "docx":
        doc = DocxDocument(BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs)
    elif ext in {"txt", "md"}:
        text = raw.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    if not text.strip():
        raise HTTPException(status_code=400, detail="No extractable text found")

    return await ingest_document(session, file.filename or "uploaded-document", text)


@app.post("/chat", response_model=ChatResponse)
async def chat(payload: ChatRequest, session: AsyncSession = Depends(get_db)):
    query_embedding = embedder.encode(payload.query)
    first_pass = await hybrid_search(session, payload.query, query_embedding, top_k=TOP_K)

    if not first_pass:
        return ChatResponse(answer="No relevant context found. Please ingest documents first.", needs_more_context=True, sources=[])

    first_context = "\n\n".join([f"[{i+1}] {row['chunk_text']}" for i, row in enumerate(first_pass)])
    draft = llm.draft_answer(payload.query, first_context)
    check = llm.self_check(payload.query, draft)

    needs_more = bool(check.get("needs_more_context", False))
    refined_query = str(check.get("query_refinement", payload.query)).strip() or payload.query

    final_results = first_pass
    if needs_more and refined_query != payload.query:
        refined_embedding = embedder.encode(refined_query)
        second_pass = await hybrid_search(session, refined_query, refined_embedding, top_k=TOP_K)
        if second_pass:
            final_results = second_pass

    final_context = "\n\n".join([f"[{i+1}] {row['chunk_text']}" for i, row in enumerate(final_results)])
    answer = llm.final_answer(payload.query, final_context)

    sources = [
        Source(
            document_id=row["document_id"],
            chunk_id=row["chunk_id"],
            title=row["title"],
            score=float(row["score"]),
            chunk_text=row["chunk_text"],
        )
        for row in final_results
    ]

    return ChatResponse(answer=answer, needs_more_context=needs_more, sources=sources)
